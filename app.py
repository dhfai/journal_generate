# # from docx import Document
# # from docx.shared import Pt
# # from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
# # from docx.oxml.ns import qn
# # from docx.oxml import OxmlElement

# # # Membuat dokumen baru
# # doc = Document()

# # # Fungsi untuk menambahkan kolom
# # def set_two_columns(section):
# #     sectPr = section._sectPr
# #     cols = sectPr.xpath('./w:cols')[0]
# #     cols.set(qn('w:num'), '2')

# # # Bagian PENDAHULUAN
# # section = doc.sections[0]
# # set_two_columns(section)

# # pendahuluan_paragraph = doc.add_paragraph()
# # pendahuluan_title = pendahuluan_paragraph.add_run("PENDAHULUAN ")
# # pendahuluan_title.font.name = "Times New Roman"
# # pendahuluan_title.font.size = Pt(12)
# # pendahuluan_title.bold = True

# # pendahuluan_paragraph.add_run("(Times New Roman 12, Bold, spasi 1, spacing before 12 pt, after 2 pt)")
# # pendahuluan_paragraph.style = doc.styles["Normal"]
# # pendahuluan_paragraph.paragraph_format.space_before = Pt(12)
# # pendahuluan_paragraph.paragraph_format.space_after = Pt(2)
# # pendahuluan_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# # pendahuluan_text = (
# #     "Bagian pendahuluan terutama berisi: (1) permasalahan penelitian; "
# #     "(2) wawasan dan rencana pemecahan masalah; (3) rumusan tujuan penelitian; "
# #     "(4) rangkuman kajian teoritik yang berkaitan dengan masalah yang diteliti. "
# #     "Pada bagian ini kadang-kadang juga dimuat harapan akan hasil dan manfaat penelitian. "
# #     "Panjang bagian pendahuluan sekitar 2-3 halaman dan diketik dengan 1 spasi. "
# #     "Template untuk format artikel ini dibuat dalam MS Word, dan selanjutnya disimpan dalam format rtf. "
# #     "Batang tubuh teks menggunakan font: Times New Roman 10, regular, spasi 1, spacing before 0 pt, after 0 pt."
# # )

# # pendahuluan_paragraph2 = doc.add_paragraph(pendahuluan_text)
# # pendahuluan_paragraph2.style = doc.styles["Normal"]
# # pendahuluan_paragraph2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
# # pendahuluan_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# # # Bagian METODE
# # metode_paragraph = doc.add_paragraph()
# # metode_title = metode_paragraph.add_run("METODE ")
# # metode_title.font.name = "Times New Roman"
# # metode_title.font.size = Pt(12)
# # metode_title.bold = True

# # metode_text = (
# #     "Pada dasarnya bagian ini menjelaskan bagaimana penelitian itu dilakukan. "
# #     "Materi pokok bagian ini adalah: (1) rancangan penelitian; (2) populasi dan sampel "
# #     "(sasaran penelitian); (3) teknik pengumpulan data dan pengembangan instrumen; "
# #     "(4) dan teknik analisis data. Untuk penelitian yang menggunakan alat dan bahan, "
# #     "perlu dituliskan spesifikasi alat dan bahannya. Spesifikasi alat menggambarkan "
# #     "kecanggihan alat yang digunakan sedangkan spesifikasi bahan menggambarkan macam bahan yang digunakan.\n\n"
# #     "Untuk penelitian kualitatif seperti penelitian tindakan kelas, etnografi, fenomenologi, "
# #     "studi kasus, dan lain-lain, perlu ditambahkan kehadiran peneliti, subyek penelitian, "
# #     "informan yang ikut membantu beserta cara-cara menggali data-data penelitian, "
# #     "lokasi dan lama penelitian serta uraian mengenai pengecekan keabsahan hasil penelitian.\n\n"
# #     "Sebaiknya dihindari pengorganisasian penulisan ke dalam “anak sub-judul” pada bagian ini. "
# #     "Namun, jika tidak bisa dihindari, cara penulisannya dapat dilihat pada bagian “Hasil dan Pembahasan”."
# # )

# # metode_paragraph2 = doc.add_paragraph(metode_text)
# # metode_paragraph2.style = doc.styles["Normal"]
# # metode_paragraph2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
# # metode_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# # # Menyimpan dokumen
# # doc.save("pendahuluan_metode.docx")



# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
# from docx.oxml.ns import qn

# # Membuat dokumen baru
# doc = Document()

# # Fungsi untuk mengatur dua kolom
# def set_two_columns(section):
#     sectPr = section._sectPr
#     cols = sectPr.xpath('./w:cols')[0]
#     cols.set(qn('w:num'), '2')

# # Bagian HASIL DAN PEMBAHASAN
# section = doc.sections[0]
# set_two_columns(section)

# hasil_paragraph = doc.add_paragraph()
# hasil_title = hasil_paragraph.add_run("HASIL DAN PEMBAHASAN ")
# hasil_title.font.name = "Times New Roman"
# hasil_title.font.size = Pt(12)
# hasil_title.bold = True

# hasil_paragraph.add_run("(Times New Roman 12, Bold, spasi 1, spacing before 12 pt, after 2 pt)")
# hasil_paragraph.style = doc.styles["Normal"]
# hasil_paragraph.paragraph_format.space_before = Pt(12)
# hasil_paragraph.paragraph_format.space_after = Pt(2)
# hasil_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# hasil_text = (
#     "Bagian ini berisi penjabaran hasil penelitian yang diperoleh serta analisis terhadap hasil tersebut. "
#     "Hasil disajikan dalam bentuk tabel, grafik, atau gambar yang relevan disertai penjelasan. "
#     "Pembahasan memuat interpretasi hasil dan relevansinya dengan kajian teori yang telah dikemukakan."
# )

# hasil_paragraph2 = doc.add_paragraph(hasil_text)
# hasil_paragraph2.style = doc.styles["Normal"]
# hasil_paragraph2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
# hasil_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# # Bagian PENUTUP
# penutup_paragraph = doc.add_paragraph()
# penutup_title = penutup_paragraph.add_run("PENUTUP ")
# penutup_title.font.name = "Times New Roman"
# penutup_title.font.size = Pt(12)
# penutup_title.bold = True

# penutup_paragraph.add_run("(Times New Roman 12, Bold, spasi 1, spacing before 12 pt, after 2 pt)")
# penutup_paragraph.style = doc.styles["Normal"]
# penutup_paragraph.paragraph_format.space_before = Pt(12)
# penutup_paragraph.paragraph_format.space_after = Pt(2)
# penutup_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# penutup_text = (
#     "Bagian penutup berisi simpulan dari hasil penelitian dan saran yang dapat diberikan berdasarkan hasil penelitian tersebut. "
#     "Simpulan harus menjawab tujuan penelitian, sedangkan saran diarahkan pada aplikasi hasil penelitian dan penelitian lanjutan."
# )

# penutup_paragraph2 = doc.add_paragraph(penutup_text)
# penutup_paragraph2.style = doc.styles["Normal"]
# penutup_paragraph2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
# penutup_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# # Bagian DAFTAR PUSTAKA
# daftar_paragraph = doc.add_paragraph()
# daftar_title = daftar_paragraph.add_run("DAFTAR PUSTAKA ")
# daftar_title.font.name = "Times New Roman"
# daftar_title.font.size = Pt(12)
# daftar_title.bold = True

# # Simpan dokumen
# doc.save("dokumen_lengkap.docx")



from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')