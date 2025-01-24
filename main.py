from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


# Membuat dokumen baru
doc = Document()

# Menambahkan header
section = doc.sections[0]
header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.text = "Penggalan Judul Artikel Jurnal"
header_paragraph.style = doc.styles["Header"]
header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Menambahkan nama penulis
penulis_paragraph = doc.add_paragraph()
penulis_run = penulis_paragraph.add_run("Nama Penulis Pertama (Times New Roman 12, Bold, spasi 1)")
penulis_run.font.name = "Times New Roman"
penulis_run.font.size = Pt(12)
penulis_run.bold = True
penulis_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Menambahkan afiliasi penulis
afiliasi_paragraph = doc.add_paragraph()
afiliasi_run = afiliasi_paragraph.add_run("Afiliasi (Program Studi, Perguruan Tinggi) dan Alamat e-mail")
afiliasi_run.font.name = "Times New Roman"
afiliasi_run.font.size = Pt(12)
afiliasi_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
afiliasi_paragraph.paragraph_format.space_after = Pt(6)

# Menambahkan nama penulis kedua
penulis_paragraph = doc.add_paragraph()
penulis_run = penulis_paragraph.add_run("Nama Penulis Kedua, dan seterusnya")
penulis_run.font.name = "Times New Roman"
penulis_run.font.size = Pt(12)
penulis_run.bold = True
penulis_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Menambahkan Abstrak dalam Bahasa Indonesia
abstrak_title_id = doc.add_paragraph()
abstrak_run_id = abstrak_title_id.add_run("Abstrak")
abstrak_run_id.font.name = "Times New Roman"
abstrak_run_id.font.size = Pt(12)
abstrak_run_id.bold = True
abstrak_title_id.alignment = WD_ALIGN_PARAGRAPH.LEFT
abstrak_title_id.paragraph_format.space_before = Pt(12)
abstrak_title_id.paragraph_format.space_after = Pt(2)

# Isi Abstrak Bahasa Indonesia
abstrak_paragraph_id = doc.add_paragraph()
abstrak_run_id = abstrak_paragraph_id.add_run(
    "Abstrak memuat uraian singkat mengenai masalah dan tujuan penelitian, metode yang digunakan, dan hasil penelitian. "
    "Tekanan penulisan abstrak terutama pada hasil penelitian. Abstrak ditulis dalam bahasa Indonesia dan Bahasa Inggris. "
    "Pengetikannya dilakukan dengan spasi tunggal dengan margin lebih sempit dari teks utama. Kata kunci harus diberikan."
)
abstrak_run_id.font.name = "Times New Roman"
abstrak_run_id.font.size = Pt(12)
abstrak_paragraph_id.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abstrak_paragraph_id.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# Menambahkan Abstrak dalam Bahasa Inggris
abstrak_title_en = doc.add_paragraph()
abstrak_run_en = abstrak_title_en.add_run("Abstract")
abstrak_run_en.font.name = "Times New Roman"
abstrak_run_en.font.size = Pt(12)
abstrak_run_en.bold = True
abstrak_title_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
abstrak_title_en.paragraph_format.space_before = Pt(12)
abstrak_title_en.paragraph_format.space_after = Pt(2)

# Isi Abstrak Bahasa Inggris
abstrak_paragraph_en = doc.add_paragraph()
abstrak_run_en = abstrak_paragraph_en.add_run(
    "The abstract provides a brief description of the problem and research objectives, the methods used, and the research results. "
    "The focus of the abstract is primarily on the research results. The abstract is written in both Indonesian and English. "
    "It is typed in single spacing with margins narrower than the main text. Keywords must be provided."
)
abstrak_run_en.font.name = "Times New Roman"
abstrak_run_en.font.size = Pt(12)
abstrak_paragraph_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abstrak_paragraph_en.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# Kata Kunci
kata_kunci_paragraph = doc.add_paragraph()
kata_kunci_run = kata_kunci_paragraph.add_run("Kata Kunci: isi, format, artikel.")
kata_kunci_run.font.name = "Times New Roman"
kata_kunci_run.font.size = Pt(12)
kata_kunci_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()


def set_two_columns(section):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

# Bagian PENDAHULUAN
section = doc.sections[0]
set_two_columns(section)

pendahuluan_paragraph = doc.add_paragraph()
pendahuluan_title = pendahuluan_paragraph.add_run("PENDAHULUAN ")
pendahuluan_title.font.name = "Times New Roman"
pendahuluan_title.font.size = Pt(12)
pendahuluan_title.bold = True

pendahuluan_paragraph.add_run("(Times New Roman 12, Bold, spasi 1, spacing before 12 pt, after 2 pt)")
pendahuluan_paragraph.style = doc.styles["Normal"]
pendahuluan_paragraph.paragraph_format.space_before = Pt(12)
pendahuluan_paragraph.paragraph_format.space_after = Pt(2)
pendahuluan_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

pendahuluan_text = (
    "Bagian pendahuluan terutama berisi: (1) permasalahan penelitian; "
    "(2) wawasan dan rencana pemecahan masalah; (3) rumusan tujuan penelitian; "
    "(4) rangkuman kajian teoritik yang berkaitan dengan masalah yang diteliti. "
    "Pada bagian ini kadang-kadang juga dimuat harapan akan hasil dan manfaat penelitian. "
    "Panjang bagian pendahuluan sekitar 2-3 halaman dan diketik dengan 1 spasi. "
    "Template untuk format artikel ini dibuat dalam MS Word, dan selanjutnya disimpan dalam format rtf. "
    "Batang tubuh teks menggunakan font: Times New Roman 10, regular, spasi 1, spacing before 0 pt, after 0 pt."
)

pendahuluan_paragraph2 = doc.add_paragraph(pendahuluan_text)
pendahuluan_paragraph2.style = doc.styles["Normal"]
pendahuluan_paragraph2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
pendahuluan_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Bagian METODE
metode_paragraph = doc.add_paragraph()
metode_title = metode_paragraph.add_run("METODE ")
metode_title.font.name = "Times New Roman"
metode_title.font.size = Pt(12)
metode_title.bold = True

metode_text = (
    "Pada dasarnya bagian ini menjelaskan bagaimana penelitian itu dilakukan. "
    "Materi pokok bagian ini adalah: (1) rancangan penelitian; (2) populasi dan sampel "
    "(sasaran penelitian); (3) teknik pengumpulan data dan pengembangan instrumen; "
    "(4) dan teknik analisis data. Untuk penelitian yang menggunakan alat dan bahan, "
    "perlu dituliskan spesifikasi alat dan bahannya. Spesifikasi alat menggambarkan "
    "kecanggihan alat yang digunakan sedangkan spesifikasi bahan menggambarkan macam bahan yang digunakan.\n\n"
    "Untuk penelitian kualitatif seperti penelitian tindakan kelas, etnografi, fenomenologi, "
    "studi kasus, dan lain-lain, perlu ditambahkan kehadiran peneliti, subyek penelitian, "
    "informan yang ikut membantu beserta cara-cara menggali data-data penelitian, "
    "lokasi dan lama penelitian serta uraian mengenai pengecekan keabsahan hasil penelitian.\n\n"
    "Sebaiknya dihindari pengorganisasian penulisan ke dalam “anak sub-judul” pada bagian ini. "
    "Namun, jika tidak bisa dihindari, cara penulisannya dapat dilihat pada bagian “Hasil dan Pembahasan”."
)

metode_paragraph2 = doc.add_paragraph(metode_text)
metode_paragraph2.style = doc.styles["Normal"]
metode_paragraph2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
metode_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY





hasil_paragraph = doc.add_paragraph()
hasil_title = hasil_paragraph.add_run("HASIL DAN PEMBAHASAN ")
hasil_title.font.name = "Times New Roman"
hasil_title.font.size = Pt(12)
hasil_title.bold = True

hasil_paragraph.add_run("(Times New Roman 12, Bold, spasi 1, spacing before 12 pt, after 2 pt)")
hasil_paragraph.style = doc.styles["Normal"]
hasil_paragraph.paragraph_format.space_before = Pt(12)
hasil_paragraph.paragraph_format.space_after = Pt(2)
hasil_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

hasil_text = (
    "Bagian ini berisi penjabaran hasil penelitian yang diperoleh serta analisis terhadap hasil tersebut. "
    "Hasil disajikan dalam bentuk tabel, grafik, atau gambar yang relevan disertai penjelasan. "
    "Pembahasan memuat interpretasi hasil dan relevansinya dengan kajian teori yang telah dikemukakan."
)

hasil_paragraph2 = doc.add_paragraph(hasil_text)
hasil_paragraph2.style = doc.styles["Normal"]
hasil_paragraph2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
hasil_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Bagian PENUTUP
penutup_paragraph = doc.add_paragraph()
penutup_title = penutup_paragraph.add_run("PENUTUP ")
penutup_title.font.name = "Times New Roman"
penutup_title.font.size = Pt(12)
penutup_title.bold = True

penutup_paragraph.add_run("(Times New Roman 12, Bold, spasi 1, spacing before 12 pt, after 2 pt)")
penutup_paragraph.style = doc.styles["Normal"]
penutup_paragraph.paragraph_format.space_before = Pt(12)
penutup_paragraph.paragraph_format.space_after = Pt(2)
penutup_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

penutup_text = (
    "Bagian penutup berisi simpulan dari hasil penelitian dan saran yang dapat diberikan berdasarkan hasil penelitian tersebut. "
    "Simpulan harus menjawab tujuan penelitian, sedangkan saran diarahkan pada aplikasi hasil penelitian dan penelitian lanjutan."
)

penutup_paragraph2 = doc.add_paragraph(penutup_text)
penutup_paragraph2.style = doc.styles["Normal"]
penutup_paragraph2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
penutup_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Bagian DAFTAR PUSTAKA
daftar_paragraph = doc.add_paragraph()
daftar_title = daftar_paragraph.add_run("DAFTAR PUSTAKA ")
daftar_title.font.name = "Times New Roman"
daftar_title.font.size = Pt(12)
daftar_title.bold = True



# Menyimpan dokumen
doc.save("aturan_penulisan.docx")
